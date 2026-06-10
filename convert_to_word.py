#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
将Markdown文件转换为Word文档
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
import re
import sys

def set_chinese_font(run, font_name='宋体', font_size=12, bold=False):
    """设置中文字体"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size)
    run.font.bold = bold

def convert_md_to_docx(md_file, docx_file):
    """转换Markdown到Word"""
    print(f"开始转换: {md_file}")
    
    # 读取Markdown文件
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    print(f"文件读取成功，内容长度: {len(content)} 字符")
    
    # 创建Word文档
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(12)
    
    # 按行处理
    lines = content.split('\n')
    print(f"总行数: {len(lines)}")
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        # 跳过空行
        if not line:
            i += 1
            continue
        
        try:
            # 处理标题
            if line.startswith('# '):
                # 一级标题
                text = line[2:]
                p = doc.add_paragraph()
                run = p.add_run(text)
                set_chinese_font(run, '黑体', 16, True)
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                
            elif line.startswith('## '):
                # 二级标题
                text = line[3:]
                p = doc.add_paragraph()
                run = p.add_run(text)
                set_chinese_font(run, '黑体', 14, True)
                
            elif line.startswith('### '):
                # 三级标题
                text = line[4:]
                p = doc.add_paragraph()
                run = p.add_run(text)
                set_chinese_font(run, '黑体', 12, True)
                
            elif line.startswith('#### '):
                # 四级标题
                text = line[5:]
                p = doc.add_paragraph()
                run = p.add_run(text)
                set_chinese_font(run, '黑体', 12, True)
                
            # 处理代码块
            elif line.startswith('```'):
                # 开始代码块
                i += 1
                code_lines = []
                while i < len(lines) and not lines[i].strip().startswith('```'):
                    code_lines.append(lines[i])
                    i += 1
                
                # 添加代码块
                if code_lines:
                    p = doc.add_paragraph()
                    code_text = '\n'.join(code_lines)
                    run = p.add_run(code_text)
                    set_chinese_font(run, 'Courier New', 10)
                    run.font.color.rgb = RGBColor(0, 0, 128)
                    
            # 处理表格
            elif line.startswith('|') and i + 1 < len(lines) and '---' in lines[i + 1]:
                # 解析表头
                headers = [cell.strip() for cell in line.split('|')[1:-1]]
                i += 2  # 跳过分隔行
                
                # 创建表格
                table = doc.add_table(rows=1, cols=len(headers))
                table.style = 'Table Grid'
                
                # 填充表头
                hdr_cells = table.rows[0].cells
                for j, header in enumerate(headers):
                    hdr_cells[j].text = header
                    for paragraph in hdr_cells[j].paragraphs:
                        for run in paragraph.runs:
                            set_chinese_font(run, '黑体', 10, True)
                
                # 填充数据行
                while i < len(lines) and lines[i].strip().startswith('|'):
                    row_data = [cell.strip() for cell in lines[i].split('|')[1:-1]]
                    if len(row_data) == len(headers):
                        row_cells = table.add_row().cells
                        for j, cell_data in enumerate(row_data):
                            row_cells[j].text = cell_data
                            for paragraph in row_cells[j].paragraphs:
                                for run in paragraph.runs:
                                    set_chinese_font(run, '宋体', 10)
                    i += 1
                continue
                
            # 处理列表
            elif line.startswith('- ') or line.startswith('* '):
                text = line[2:]
                p = doc.add_paragraph(text, style='List Bullet')
                for run in p.runs:
                    set_chinese_font(run, '宋体', 12)
                    
            elif re.match(r'^\d+\. ', line):
                text = re.sub(r'^\d+\. ', '', line)
                p = doc.add_paragraph(text, style='List Number')
                for run in p.runs:
                    set_chinese_font(run, '宋体', 12)
                    
            # 处理普通段落
            else:
                # 处理加粗 **text**
                parts = re.split(r'\*\*(.*?)\*\*', line)
                p = doc.add_paragraph()
                for idx, part in enumerate(parts):
                    run = p.add_run(part)
                    if idx % 2 == 1:  # 加粗部分
                        set_chinese_font(run, '宋体', 12, True)
                    else:
                        set_chinese_font(run, '宋体', 12)
        except Exception as e:
            print(f"处理第{i+1}行时出错: {e}")
            print(f"行内容: {line}")
        
        i += 1
    
    # 保存文档
    doc.save(docx_file)
    print(f"转换完成！Word文档已保存至: {docx_file}")

if __name__ == '__main__':
    md_file = r'c:\Users\Liu17\OneDrive\Desktop\lhh35\毕业设计说明书_8000字版.md'
    docx_file = r'c:\Users\Liu17\OneDrive\Desktop\lhh35\毕业设计说明书_8000字版.docx'
    convert_md_to_docx(md_file, docx_file)
